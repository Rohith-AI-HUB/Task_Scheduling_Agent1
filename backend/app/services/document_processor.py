"""
Document Processor Service
Handles text extraction from various document formats including:
- PDF files
- DOCX/DOC files
- Code files (Python, JavaScript, etc.)
- Images (via Tesseract OCR)
"""

import os
import platform
import logging
from typing import Optional
from pathlib import Path

# PDF processing
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

# DOCX processing
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# Image processing and OCR
try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import pytesseract
except ImportError:
    pytesseract = None

logger = logging.getLogger(__name__)

# Configure Tesseract path for Windows
if pytesseract is not None and platform.system() == 'Windows':
    tesseract_paths = [
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        r'C:\Users\rohit\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'
    ]
    for path in tesseract_paths:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            break

# Supported file types by category
SUPPORTED_DOCUMENTS = {
    'pdf': ['pdf'],
    'docx': ['doc', 'docx'],
    'text': ['txt', 'md', 'markdown'],
    'code': ['py', 'js', 'jsx', 'ts', 'tsx', 'java', 'cpp', 'c', 'h', 'hpp',
             'css', 'html', 'json', 'xml', 'yaml', 'yml', 'sql', 'sh', 'bash',
             'go', 'rs', 'rb', 'php', 'swift', 'kt', 'scala'],
    'image': ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp']
}

# All supported extensions flattened
ALL_SUPPORTED_EXTENSIONS = set()
for extensions in SUPPORTED_DOCUMENTS.values():
    ALL_SUPPORTED_EXTENSIONS.update(extensions)


def get_file_extension(filename: str) -> str:
    """Extract lowercase file extension from filename."""
    if '.' not in filename:
        return ''
    return filename.lower().rsplit('.', 1)[-1]


def get_document_type(filename: str) -> str:
    """Determine document type category from filename."""
    ext = get_file_extension(filename)
    for doc_type, extensions in SUPPORTED_DOCUMENTS.items():
        if ext in extensions:
            return doc_type
    return 'unknown'


def is_supported_file(filename: str) -> bool:
    """Check if file type is supported for processing."""
    ext = get_file_extension(filename)
    return ext in ALL_SUPPORTED_EXTENSIONS


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text content from a PDF file.

    Args:
        file_path: Path to the PDF file

    Returns:
        Extracted text content
    """
    if PdfReader is None:
        return "[Error: pypdf is not installed. Install it to enable PDF extraction.]"
    try:
        reader = PdfReader(file_path)
        text_parts = []

        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {i + 1} ---\n{page_text}")

        return "\n\n".join(text_parts).strip()
    except Exception as e:
        logger.error(f"PDF extraction error for {file_path}: {e}")
        return f"[Error extracting PDF: {str(e)}]"


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text content from a DOCX file.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text content
    """
    if DocxDocument is None:
        return "[Error: python-docx is not installed. Install it to enable DOCX extraction.]"
    try:
        doc = DocxDocument(file_path)
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))
            if table_text:
                text_parts.append("\n[Table]\n" + "\n".join(table_text))

        return "\n\n".join(text_parts).strip()
    except Exception as e:
        logger.error(f"DOCX extraction error for {file_path}: {e}")
        return f"[Error extracting DOCX: {str(e)}]"


def extract_text_from_image(file_path: str) -> str:
    """
    Extract text from an image using Tesseract OCR.

    Args:
        file_path: Path to the image file

    Returns:
        Extracted text content from OCR
    """
    if Image is None:
        return "[Error: Pillow is not installed. Install it to enable image extraction.]"
    if pytesseract is None:
        return "[Error: pytesseract is not installed. Install it to enable OCR extraction.]"
    try:
        # Open and preprocess image for better OCR
        image = Image.open(file_path)

        # Convert to RGB if necessary (for PNG with transparency)
        if image.mode in ('RGBA', 'LA', 'P'):
            background = Image.new('RGB', image.size, (255, 255, 255))
            if image.mode == 'P':
                image = image.convert('RGBA')
            background.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
            image = background
        elif image.mode != 'RGB':
            image = image.convert('RGB')

        # Perform OCR
        text = pytesseract.image_to_string(image)

        if not text.strip():
            return "[No text detected in image]"

        return f"[OCR Extracted Text]\n{text.strip()}"
    except pytesseract.TesseractNotFoundError:
        logger.error("Tesseract OCR not installed or not found")
        return "[Error: Tesseract OCR not installed. Please install Tesseract to enable image text extraction.]"
    except Exception as e:
        logger.error(f"OCR extraction error for {file_path}: {e}")
        return f"[Error extracting text from image: {str(e)}]"


def extract_text_from_code(file_path: str) -> str:
    """
    Extract content from a code/text file.

    Args:
        file_path: Path to the code file

    Returns:
        File content with syntax indication
    """
    try:
        # Try different encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        content = None

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue

        if content is None:
            return "[Error: Could not decode file content]"

        ext = get_file_extension(file_path)
        return f"[Code File: .{ext}]\n{content}"
    except Exception as e:
        logger.error(f"Code file read error for {file_path}: {e}")
        return f"[Error reading code file: {str(e)}]"


def extract_text_from_text_file(file_path: str) -> str:
    """
    Extract content from plain text files.

    Args:
        file_path: Path to the text file

    Returns:
        File content
    """
    try:
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        content = None

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue

        if content is None:
            return "[Error: Could not decode file content]"

        return content
    except Exception as e:
        logger.error(f"Text file read error for {file_path}: {e}")
        return f"[Error reading text file: {str(e)}]"


def process_uploaded_document(file_path: str, filename: str) -> dict:
    """
    Process an uploaded document and extract its text content.

    Args:
        file_path: Path to the uploaded file
        filename: Original filename

    Returns:
        Dictionary containing:
        - success: bool
        - extracted_text: str
        - document_type: str
        - character_count: int
        - error: str (if failed)
    """
    if not os.path.exists(file_path):
        return {
            "success": False,
            "error": "File not found",
            "extracted_text": "",
            "document_type": "unknown",
            "character_count": 0
        }

    doc_type = get_document_type(filename)

    # Route to appropriate extractor
    extractors = {
        'pdf': extract_text_from_pdf,
        'docx': extract_text_from_docx,
        'text': extract_text_from_text_file,
        'code': extract_text_from_code,
        'image': extract_text_from_image
    }

    extractor = extractors.get(doc_type)

    if not extractor:
        return {
            "success": False,
            "error": f"Unsupported document type: {doc_type}",
            "extracted_text": "",
            "document_type": doc_type,
            "character_count": 0
        }

    try:
        text = extractor(file_path)

        # Check for extraction errors
        if text.startswith("[Error"):
            return {
                "success": False,
                "error": text,
                "extracted_text": "",
                "document_type": doc_type,
                "character_count": 0
            }

        return {
            "success": True,
            "extracted_text": text,
            "document_type": doc_type,
            "character_count": len(text),
            "filename": filename
        }
    except Exception as e:
        logger.error(f"Document processing error: {e}")
        return {
            "success": False,
            "error": str(e),
            "extracted_text": "",
            "document_type": doc_type,
            "character_count": 0
        }


def get_document_summary(extracted_text: str, max_length: int = 500) -> str:
    """
    Get a summary/preview of extracted document text.

    Args:
        extracted_text: Full extracted text
        max_length: Maximum length for summary

    Returns:
        Truncated text with ellipsis if needed
    """
    if len(extracted_text) <= max_length:
        return extracted_text

    return extracted_text[:max_length].rsplit(' ', 1)[0] + "..."


def check_tesseract_installed() -> bool:
    """Check if Tesseract OCR is properly installed and accessible."""
    if pytesseract is None:
        return False
    try:
        pytesseract.get_tesseract_version()
        return True
    except pytesseract.TesseractNotFoundError:
        return False
    except Exception:
        return False
